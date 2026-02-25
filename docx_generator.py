import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


# --- Color palette ---
COLOR_PRIMARY = RGBColor(0x1A, 0x56, 0x8E)  # Dark blue
COLOR_SECONDARY = RGBColor(0x2E, 0x86, 0xC1)  # Medium blue
COLOR_ACCENT = RGBColor(0x5D, 0xAE, 0xD8)  # Light blue
COLOR_DARK = RGBColor(0x2C, 0x3E, 0x50)  # Dark gray-blue
COLOR_LIGHT_BG = RGBColor(0xEB, 0xF5, 0xFB)  # Very light blue
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)  # Dark gray
COLOR_MUTED = RGBColor(0x77, 0x77, 0x77)  # Medium gray
COLOR_TERMINATE = RGBColor(0xC0, 0x39, 0x2B)  # Red
COLOR_SKIP = RGBColor(0xE6, 0x7E, 0x22)  # Orange


def set_cell_shading(cell, color_hex: str):
    """Set background color for a table cell."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    shading_elm.append(shading)


def add_horizontal_line(doc, color="1A568E"):
    """Add a horizontal line paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): color,
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_cover_page(doc, project_data: dict):
    """Create a professional cover page."""
    # Add spacing before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("QUESTIONÁRIO DE PESQUISA")
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.bold = True

    # Subtitle — research objective
    objective = project_data.get("research_objective", "Pesquisa de Mercado")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(objective)
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_DARK

    # Separator line
    add_horizontal_line(doc)

    # Metadata table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)

    meta_items = [
        ("Público-alvo", project_data.get("target_audience", "—")),
        ("Metodologia", project_data.get("methodology", "—")),
        ("LOI estimada", f"{project_data.get('estimated_loi_minutes', '—')} minutos"),
        ("Total de perguntas", str(project_data.get("total_questions", "—"))),
        ("Data", datetime.now().strftime("%d/%m/%Y")),
        ("Status", "RASCUNHO — Para revisão"),
    ]

    table = doc.add_table(rows=len(meta_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(meta_items):
        row = table.rows[i]
        # Label cell
        cell_label = row.cells[0]
        cell_label.width = Cm(5)
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY

        # Value cell
        cell_value = row.cells[1]
        cell_value.width = Cm(10)
        p = cell_value.paragraphs[0]
        run = p.add_run(f"  {value}")
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TEXT

    # Confidentiality note
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENCIAL")
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Este documento contém informações proprietárias destinadas exclusivamente à equipe do projeto.")
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

    doc.add_page_break()


def format_question_type(q_type: str) -> str:
    """Return a human-readable label for question types."""
    TYPE_LABELS = {
        "single_choice": "Resposta Única",
        "multiple_choice": "Resposta Múltipla",
        "scale_numeric": "Escala Numérica",
        "scale_likert": "Escala Likert",
        "nps": "NPS (0-10)",
        "ranking": "Ranking / Ordenação",
        "open_text": "Resposta Aberta",
        "matrix": "Matriz / Grid",
    }
    return TYPE_LABELS.get(q_type, q_type.replace("_", " ").title())


def add_section_header(doc, section_data: dict):
    """Add a section header with title and description."""
    # Section title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(f"{section_data['id']}. {section_data['title']}")
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.bold = True

    # Section description
    if section_data.get("description"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(section_data["description"])
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED
        run.font.italic = True

    add_horizontal_line(doc, "5DAED8")


def add_question(doc, question: dict):
    """Add a formatted question to the document."""
    q_type_label = format_question_type(question.get("type", ""))

    # Question ID + type badge
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(2)

    run = p.add_run(f"{question['id']}")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY

    run = p.add_run(f"  [{q_type_label}]")
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

    if question.get("required", True):
        run = p.add_run("  *Obrigatória")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_TERMINATE

    # Question text
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(question["text"])
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXT
    run.font.bold = True

    # Instruction
    if question.get("instruction"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(question["instruction"])
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED
        run.font.italic = True

    # --- Render based on type ---
    q_type = question.get("type", "")

    if q_type in ("single_choice", "multiple_choice"):
        _render_choice_options(doc, question)

    elif q_type in ("scale_numeric", "nps"):
        _render_scale(doc, question)

    elif q_type == "scale_likert":
        _render_likert(doc, question)

    elif q_type == "ranking":
        _render_ranking(doc, question)

    elif q_type == "open_text":
        _render_open_text(doc, question)

    elif q_type == "matrix":
        _render_matrix(doc, question)

    # Programming note
    if question.get("programming_note"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(f"📋 Programação: {question['programming_note']}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)

    # Methodological note
    if question.get("methodological_note"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(f"🔬 Nota metodológica: {question['methodological_note']}")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_MUTED


def _render_choice_options(doc, question: dict):
    """Render single/multiple choice options."""
    options = question.get("options", [])
    is_single = question.get("type") == "single_choice"
    marker = "○" if is_single else "☐"
    randomize = question.get("randomize_options", False)

    if randomize:
        p = doc.add_paragraph()
        run = p.add_run("⟳ RANDOMIZAR ORDEM")
        run.font.size = Pt(7)
        run.font.color.rgb = COLOR_ACCENT
        run.font.bold = True

    for opt in options:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(1)

        if isinstance(opt, dict):
            code = opt.get("code", "")
            text = opt.get("text", "")
            routing = opt.get("routing", "")

            run = p.add_run(f"{marker}  {code}. {text}")
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_TEXT

            if routing and routing not in ("CONTINUE", ""):
                if routing == "TERMINATE":
                    run = p.add_run(f"  → ENCERRAR")
                    run.font.size = Pt(8)
                    run.font.color.rgb = COLOR_TERMINATE
                    run.font.bold = True
                else:
                    run = p.add_run(f"  → Ir para {routing}")
                    run.font.size = Pt(8)
                    run.font.color.rgb = COLOR_SKIP
        else:
            run = p.add_run(f"{marker}  {opt}")
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_TEXT


def _render_scale(doc, question: dict):
    """Render a numeric scale (including NPS)."""
    scale_min = question.get("scale_min", 0)
    scale_max = question.get("scale_max", 10)
    anchor_min = question.get("anchor_min", "")
    anchor_max = question.get("anchor_max", "")

    num_points = scale_max - scale_min + 1
    cols = min(num_points, 11)  # Max 11 columns for readability

    table = doc.add_table(rows=2, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Number row
    for i in range(cols):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(scale_min + i))
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY

    # Anchor row
    if anchor_min or anchor_max:
        cell_min = table.rows[1].cells[0]
        p = cell_min.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(anchor_min)
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_MUTED

        cell_max = table.rows[1].cells[cols - 1]
        p = cell_max.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(anchor_max)
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_MUTED

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def _render_likert(doc, question: dict):
    """Render Likert scale options."""
    options = question.get("options", [
        "Discordo totalmente",
        "Discordo parcialmente",
        "Não concordo nem discordo",
        "Concordo parcialmente",
        "Concordo totalmente",
    ])

    for i, opt in enumerate(options, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(1)

        text = opt if isinstance(opt, str) else opt.get("text", "")
        run = p.add_run(f"○  {i}. {text}")
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TEXT


def _render_ranking(doc, question: dict):
    """Render ranking items."""
    items = question.get("options", question.get("items", []))

    p = doc.add_paragraph()
    run = p.add_run("Ordene de mais importante (1º) a menos importante:")
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED
    run.font.italic = True

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(1)

        text = item if isinstance(item, str) else item.get("text", "")
        run = p.add_run(f"___  {text}")
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TEXT


def _render_open_text(doc, question: dict):
    """Render open text field."""
    # Draw a box representation
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.width = Cm(14)

    # Set border
    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "bottom", "left", "right"):
        border = borders.makeelement(
            qn(f"w:{edge}"),
            {qn("w:val"): "single", qn("w:sz"): "4", qn("w:color"): "CCCCCC"},
        )
        borders.append(border)
    tcPr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(" ")
    run.font.size = Pt(10)

    max_chars = question.get("max_chars", "")
    if max_chars:
        p = doc.add_paragraph()
        run = p.add_run(f"Máximo: {max_chars} caracteres")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_MUTED


def _render_matrix(doc, question: dict):
    """Render a matrix/grid question."""
    rows_data = question.get("rows", question.get("items", []))
    cols_data = question.get("columns", question.get("scale_points", []))

    if not rows_data or not cols_data:
        p = doc.add_paragraph()
        run = p.add_run("[Matriz — configurar linhas e colunas]")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED
        return

    num_cols = len(cols_data) + 1
    table = doc.add_table(rows=len(rows_data) + 1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for j, col_label in enumerate(cols_data, 1):
        text = col_label if isinstance(col_label, str) else col_label.get("text", "")
        p = header_row.cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY

    # Data rows
    for i, row_item in enumerate(rows_data):
        text = row_item if isinstance(row_item, str) else row_item.get("text", "")
        row = table.rows[i + 1]

        p = row.cells[0].paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_TEXT

        for j in range(1, num_cols):
            p = row.cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("○")
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_MUTED


def add_methodology_notes(doc, notes: dict):
    """Add methodology notes section at the end."""
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("ANEXO METODOLÓGICO")
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.bold = True

    add_horizontal_line(doc)

    sections = [
        ("Amostragem", notes.get("sampling", "")),
        ("Cotas", notes.get("quotas", "")),
        ("Limitações", notes.get("limitations", "")),
    ]

    for title, content in sections:
        if content:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(title)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = COLOR_DARK

            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_TEXT

    # Biases mitigated
    biases = notes.get("biases_mitigated", [])
    if biases:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run("Vieses Controlados no Design")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_DARK

        for bias in biases:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f"• {bias}")
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_TEXT


def generate_questionnaire_docx(questionnaire_json: dict) -> bytes:
    """Generate a complete .docx questionnaire from JSON data."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = COLOR_TEXT

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Cover page
    project_data = questionnaire_json.get("project_summary", {})
    create_cover_page(doc, project_data)

    # Platform notes
    platform_notes = project_data.get("platform_notes", "")
    if platform_notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"Notas para programação: {platform_notes}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)
        doc.add_page_break()

    # Sections and questions
    sections = questionnaire_json.get("sections", [])
    for sec_idx, section in enumerate(sections):
        if sec_idx > 0:
            # Add spacing between sections (but not page break for every section)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)

        add_section_header(doc, section)

        questions = section.get("questions", [])
        for question in questions:
            add_question(doc, question)

    # Methodology notes
    method_notes = questionnaire_json.get("methodological_notes", {})
    if method_notes:
        add_methodology_notes(doc, method_notes)

    # Footer note
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Fim do Questionário —")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_MUTED
    run.font.italic = True

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
